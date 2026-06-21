from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt


ROOT = Path(__file__).resolve().parents[1]
TEMPLATE = ROOT / "docs" / "测试方案修改-zjl.docx"
OUT = ROOT / "docs" / "智联计算系统测试方案-正式修订版.docx"
ASSET = ROOT / "docs" / "assets" / "evaluation-plan"


def clear_document(doc: Document) -> None:
    body = doc._body._element
    for child in list(body):
        if child.tag != qn("w:sectPr"):
            body.remove(child)


def set_run_font(run, size: float = 10.5, bold: bool = False) -> None:
    run.font.name = "宋体"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    run.font.size = Pt(size)
    run.bold = bold


def add_paragraph(parent, text: str = "", *, bold: bool = False, size: float = 10.5, align=None):
    p = parent.add_paragraph()
    if align is not None:
        p.alignment = align
    if text:
        run = p.add_run(text)
        set_run_font(run, size=size, bold=bold)
    return p


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_width(cell, width_twips: int) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.first_child_found_in("w:tcW")
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width_twips))
    tc_w.set(qn("w:type"), "dxa")


def clear_cell(cell) -> None:
    cell.text = ""
    for p in cell.paragraphs:
        p.text = ""


def write_cell(cell, text: str, *, bold: bool = False) -> None:
    clear_cell(cell)
    p = cell.paragraphs[0]
    run = p.add_run(text)
    set_run_font(run, bold=bold)


def add_image_to_cell(cell, caption: str, filename: str) -> None:
    cap = cell.add_paragraph()
    cap.paragraph_format.space_before = Pt(4)
    cap.paragraph_format.space_after = Pt(2)
    run = cap.add_run(caption)
    set_run_font(run, size=9.5, bold=True)
    img_p = cell.add_paragraph()
    img_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    img_run = img_p.add_run()
    img_run.add_picture(str(ASSET / filename), width=Inches(5.25))


def add_steps(cell, steps: list[dict]) -> None:
    clear_cell(cell)
    for idx, step in enumerate(steps, 1):
        p = cell.add_paragraph()
        p.paragraph_format.space_after = Pt(2)
        run = p.add_run(f"{idx}. {step['text']}")
        set_run_font(run)
        if step.get("image"):
            add_image_to_cell(cell, step["caption"], step["image"])


def add_multiline(cell, lines: list[str]) -> None:
    clear_cell(cell)
    for idx, line in enumerate(lines):
        p = cell.paragraphs[0] if idx == 0 else cell.add_paragraph()
        run = p.add_run(line)
        set_run_font(run)


def add_case_table(doc: Document, case: dict) -> None:
    add_paragraph(doc, case["heading"], bold=True, size=14)
    add_paragraph(doc, case["table_title"], bold=True, size=11)
    labels = ["用例编号", "测试目的"]
    values = [
        case["id"],
        case["purpose"],
    ]
    for item in case.get("metric_rows", []):
        labels.append(item["label"])
        values.append(item["text"])
    if case.get("topology", True):
        labels.append("组网拓扑")
        values.append(case.get("topology_text", "见本测试方案统一组网拓扑。"))
    labels += ["前置条件", "测试步骤", "预期结果", "测试结果", ""]
    values += [
        case["preconditions"],
        case["steps"],
        case["expected"],
        case["result"],
        case.get("note", "截图已随测试步骤穿插留档，用于证明页面操作、单任务指标、成功率计算和任务工单详情。"),
    ]

    table = doc.add_table(rows=len(labels), cols=2)
    table.style = "Table Grid"
    table.autofit = False
    for row_idx, (label, value) in enumerate(zip(labels, values)):
        row = table.rows[row_idx]
        set_cell_width(row.cells[0], 1800)
        set_cell_width(row.cells[1], 7800)
        write_cell(row.cells[0], label, bold=True)
        set_cell_shading(row.cells[0], "EAF2F8")
        if label == "测试步骤":
            add_steps(row.cells[1], value)
        elif isinstance(value, list):
            add_multiline(row.cells[1], value)
        else:
            write_cell(row.cells[1], value)
    add_paragraph(doc)


def build() -> None:
    doc = Document(str(TEMPLATE))
    clear_document(doc)
    section = doc.sections[0]
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.8)
    section.left_margin = Inches(0.8)
    section.right_margin = Inches(0.8)

    title = add_paragraph(doc, "智联计算系统测试方案（正式修订版）", bold=True, size=18, align=WD_ALIGN_PARAGRAPH.CENTER)
    title.paragraph_format.space_after = Pt(12)
    add_paragraph(
        doc,
        "本文档仅保留当前负责的两个验收指标相关测试用例：业务目标成功率和意图解析参数提取准确率。"
        "测试业务参数采用固定输入，任务开始与结束时间由系统工单字段填写，不采用人工计时；页面需通过测试工单列表展示每个任务的指标结果、阈值、是否达标和汇总成功率计算过程。",
    )
    add_paragraph(
        doc,
        "业务目标成功率采用节点同业务参数历史基线作为比较对象。矩阵乘法吞吐量越高越好，单任务达标条件为实际吞吐量不低于节点基线的 80%；"
        "视频AI推理时延越低越好，当前验收口径为实际 P90 帧推理时延不超过节点基线 P90 时延的 1.5 倍。"
        "该裕量用于覆盖容器化运行、网络转发和系统调度带来的正常波动，不用于放宽业务目标本身。",
    )

    cases = [
        {
            "heading": "3.1.5 矩阵乘法计算业务目标成功率测试",
            "table_title": "表1-5 矩阵乘法计算业务目标成功率测试",
            "id": "1-5",
            "purpose": "验证高通量计算模态下矩阵乘法计算业务目标成功率，并证明系统能够展示单任务计算速率指标和成功率计算过程。",
            "metric_rows": [
                {"label": "测试指标", "text": "矩阵计算速率 effective_gflops，单位为 GFLOPS，数值越高表示计算能力越强。"},
                {"label": "基线说明", "text": "正式测试前，对参与调度的算力节点使用相同业务参数进行基线测试，取稳定历史结果作为该节点基线。"},
                {"label": "判定方法", "text": "单个工单实际计算速率不低于对应节点基线的 80% 时，判定该工单达标。"},
                {"label": "统计口径", "text": "业务目标成功率 = 达标工单数 / 可评价工单数 × 100%；正式验收要求可评价工单不少于 30 个，成功率不低于 90%。"},
            ],
            "preconditions": [
                "1. 按统一组网拓扑完成智联计算系统部署，管理端、前端、任务编排服务、节点代理和数据库均正常运行。",
                "2. 已准备矩阵乘法计算任务模板，业务数据按 source → compute → sink 链路流转。",
                "3. 已在参与调度的算力节点上完成相同固定参数的历史基线测试。",
                "4. 正式验收前可清理历史测评轮次和残留容器实例，避免旧数据影响当前轮次统计。",
                "5. 测试参数固定为 matrix_size=1024、batch_count=50、observation_duration_sec=10、min_samples=5；单个工单业务时间窗口不小于 5 分钟。",
            ],
            "steps": [
                {"text": "管理员登录系统，进入业务测评页面，选择“矩阵乘法计算任务”。"},
                {"text": "检查基线区域，确认参与调度的计算节点均有对应基线值，并查看稳定/波动提示。"},
                {"text": "创建 30 个测评工单，系统生成当前验收轮次 high_throughput_matmul-20260621130040，并写入固定业务参数。", "caption": "图1-5-1 矩阵乘法任务基线和固定测评参数", "image": "matmul-benchmark-setup.png"},
                {"text": "路由模块为每个工单写回 source、compute、sink 放置结果，compute 节点携带 GPU 分配信息；部署前由系统校验节点和 GPU 分配冲突。"},
                {"text": "点击运行测评，系统按当前并发任务数小批次执行本轮工单，资源暂不可用时自动等待。"},
                {"text": "业务计算节点跳过预热阶段后持续采集吞吐样本，计算有效浮点运算吞吐量，计算方式为 (2 × N³ × 本轮计算批次数) / 本轮耗时 / 1e9。"},
                {"text": "系统上报单任务实际计算速率、节点基线、达标阈值、采样点数和运行结果。"},
                {"text": "在测试工单列表中查看每个任务的实际指标、阈值和是否达标。", "caption": "图1-5-2 矩阵乘法任务单工单指标结果表", "image": "matmul-benchmark-result.png"},
                {"text": "系统按“业务目标成功率 = 达标任务数 / 可评价任务数 × 100%”汇总展示，当前轮次 30/30 达标。", "caption": "图1-5-3 矩阵乘法任务业务目标成功率汇总", "image": "matmul-benchmark-evidence-table.png"},
                {"text": "展开工单详情，检查业务输入、路由放置、GPU 分配、部署状态、计算输出和结果文件。", "caption": "图1-5-4 矩阵乘法工单详情", "image": "matmul-order-detail-overview.png"},
            ],
            "expected": "业务目标成功率不小于 90%，且已评估任务数不少于 30 个。",
            "result": "当前管理节点验收轮次 high_throughput_matmul-20260621130040 共 30 个工单，30 个可评价且全部达标，业务目标成功率为 100.0%，满足验收要求。",
        },
        {
            "heading": "3.1.6 视频AI推理业务目标成功率测试",
            "table_title": "表1-6 视频AI推理业务目标成功率测试",
            "id": "1-6",
            "purpose": "验证低时延转发模态下视频AI推理业务目标成功率，并证明系统能够展示单任务帧推理时延、检测框和成功率计算过程。",
            "metric_rows": [
                {"label": "测试指标", "text": "帧推理 P90 时延 frame_latency_p90_ms，单位为 ms，数值越低表示低时延处理能力越强。"},
                {"label": "基线说明", "text": "正式测试前，对参与调度的算力节点使用相同固定视频和相同视频参数进行基线测试，取稳定历史结果作为该节点基线。"},
                {"label": "判定方法", "text": "单个工单实际 P90 时延不超过对应节点基线 P90 时延的 1.5 倍时，判定该工单达标。"},
                {"label": "统计口径", "text": "业务目标成功率 = 达标工单数 / 可评价工单数 × 100%；正式验收要求可评价工单不少于 30 个，成功率不低于 90%。"},
            ],
            "preconditions": [
                "1. 按统一组网拓扑完成智联计算系统部署，管理端、前端、任务编排服务、节点代理和数据库均正常运行。",
                "2. 已准备视频AI推理任务模板，业务数据按 source → compute → sink 链路流转。",
                "3. 已准备固定测试视频 bottle-detection.mp4 和 YOLOv5n ONNX 权重文件，测试方案验收时使用镜像内置固定文件。",
                "4. 已在参与调度的算力节点上完成相同视频参数的历史基线测试。",
                "5. 正式验收前可清理历史测评轮次和残留容器实例，避免旧数据影响当前轮次统计。",
                "6. 测试参数固定为 frame_count=100、frame_stride=30、warmup_frames=10、measured_frames=30；单个工单业务时间窗口不小于 5 分钟。",
            ],
            "steps": [
                {"text": "管理员登录系统，进入业务测评页面，选择“视频AI推理任务”。"},
                {"text": "检查基线区域，确认参与调度的计算节点均有当前视频参数对应的基线时延，并查看稳定/波动提示。"},
                {"text": "创建 30 个测评工单，系统生成当前视频AI推理正式验收轮次，并写入固定视频输入参数。", "caption": "图1-6-1 视频AI推理任务基线和固定测评参数", "image": "video-benchmark-setup.png"},
                {"text": "路由模块写回 source、compute、sink 放置结果，compute 节点携带 GPU 分配信息；部署系统对 GPU 冲突进行兜底校验。"},
                {"text": "点击运行测评，系统按当前并发任务数小批次执行本轮工单，资源暂不可用时自动等待。"},
                {"text": "source 节点读取固定测试视频并按 frame_stride 抽帧发送给 compute 节点。"},
                {"text": "compute 节点加载 YOLOv5n 权重，跳过 warmup_frames 后统计 measured_frames 的逐帧推理时延，并生成中文分类标签和检测框。"},
                {"text": "sink 节点汇总并上报 frame_latency_p90_ms、有效帧数、检测类别、画框坐标和带框预览图。"},
                {"text": "在测试工单列表中查看每个任务的 P90 帧推理时延、阈值和是否达标；视频时延越低越好，达标阈值按“节点基准 P90 时延 × 1.5”计算。", "caption": "图1-6-2 视频AI推理任务单工单帧时延结果表", "image": "video-benchmark-result.png"},
                {"text": "系统按“业务目标成功率 = 达标任务数 / 可评价任务数 × 100%”汇总展示。正式留档需使用不少于 30 个可评价工单的验收轮次。", "caption": "图1-6-3 视频AI推理任务业务目标成功率汇总", "image": "video-benchmark-evidence-table.png"},
                {"text": "展开工单详情，检查业务输入、路由放置、GPU 分配、部署状态、检测类别、画框坐标和带框预览图。", "caption": "图1-6-4 视频AI推理工单详情带框预览", "image": "video-order-detail-result-preview.png"},
            ],
            "expected": "业务目标成功率不小于 90%，且已评估任务数不少于 30 个。",
            "result": "当前管理节点验收轮次 low_latency_video_pipeline-20260621133404 共 30 个工单，30 个可评价且全部达标，业务目标成功率为 100.0%，满足验收要求；工单详情可展示 YOLO 检测框、中文类别标签、GPU 分配和帧时延元数据。",
        },
        {
            "heading": "3.2.1 用户意图解析功能测试",
            "table_title": "表2-1 用户意图解析功能测试",
            "id": "2-1",
            "purpose": "验证用户通过对话输入业务需求后，系统能够生成结构化解析结果，并完成工单提交。",
            "topology": False,
            "metric_rows": [
                {"label": "判定方法", "text": "本用例为功能性测试，不计算准确率；系统能够根据固定输入正常返回结构化解析结果并完成工单提交，即判定通过。"},
            ],
            "preconditions": [
                "1. 系统管理端和用户端模块部署正常，用户已完成登录。",
                "2. 后端意图解析工作流可正常运行，已配置支持的业务类型、节点名称和参数校验规则。",
                "3. 系统中已维护 h1-h13 运营商端点和 compute-1~3 算力节点；用户输入的源/目的端点使用 h1-h13，算力节点由系统或路由模块分配。",
            ],
            "steps": [
                {"text": "用户进入意图对话页面。"},
                {"text": "输入完整测试文本：我想部署视频AI推理业务，从 h1 到 h2，使用 yolo 模型，分辨率 720p，抽取 100 帧，30fps，现在开始运行 10 分钟，采用低时延转发策略。"},
                {"text": "系统返回结构化解析结果，包括业务类型、所属模态、源端点、目的端点、业务时间窗口、任务参数和路由策略。"},
                {"text": "用户确认提交任务，系统生成工单并在工单中心展示任务状态。", "caption": "图2-1-1 用户意图对话与参数解析结果", "image": "intent-chat-page.png"},
                {"text": "新建对话，输入缺失参数测试文本：我想做矩阵乘法任务，从 h1 到 h2，1024 阶矩阵，50 批。确认系统提示需要补充业务运行时间，且补齐前不能提交工单。"},
                {"text": "继续补充：现在开始运行 10 分钟。确认系统合并上一轮草稿并允许提交。"},
                {"text": "新建对话，输入非法节点测试文本：从 compute-1 到 h2 做矩阵乘法任务，1024 阶矩阵，50 批，现在开始运行 10 分钟。确认系统提示源端点不合法或不能使用计算节点作为用户源/目的端点，且不能提交工单。"},
                {"text": "新建对话，输入非法参数测试文本：从 h1 到 h2 做视频AI推理，720p，抽取 100 帧，500fps，现在开始运行 10 分钟。确认系统提示参数不合理，要求修正后才能提交。"},
            ],
            "expected": "系统能够完成完整输入解析和工单提交；对缺失参数能够追问补全；对不存在节点、计算节点误填为源/目的端点、明显不合理参数能够给出中文提示并阻止提交。",
            "result": "当前系统已实现意图对话、参数草稿、多轮补全、节点与参数校验、确认提交、工单生成和路由请求记录；正式验收时保存对话页面、结构化参数、缺参/非法值提示和工单详情截图。",
        },
        {
            "heading": "3.2.2 意图解析参数提取准确率测试",
            "table_title": "表2-2 意图解析参数提取准确率测试",
            "id": "2-2",
            "purpose": "验证固定数据集下意图解析参数提取准确率达到不低于 90% 的验收要求。",
            "topology": False,
            "metric_rows": [
                {"label": "测试指标", "text": "意图解析参数提取准确率，用于衡量系统从自然语言需求中提取关键任务参数的正确程度。"},
                {"label": "判定方法", "text": "单条样本的关键字段与人工标注结果一致时，判定该样本正确；期望缺失的字段不应被系统幻觉补全。"},
                {"label": "统计口径", "text": "参数提取准确率 = 正确样本数 / 总样本数 × 100%；正式验收要求准确率不低于 90%。"},
            ],
            "preconditions": [
                "1. 管理员已登录系统，意图评测页面可正常访问。",
                "2. 已准备固定意图解析数据集，数据集由自然语言业务需求文本和人工标注结果组成，覆盖不同业务类型、模态、完整输入和缺失参数等常见情况。",
                "3. 在线大模型/智能体评测所需模型配置可用。",
            ],
            "steps": [
                {"text": "管理员进入意图评测页面，确认固定数据集规模和验收目标。"},
                {"text": "选择在线评测模型，运行智能体/大模型意图解析评测。"},
                {"text": "系统逐条调用在线接口完成数据集评测，展示评测编号、模型名称、总样本数、正确样本数和准确率。"},
                {"text": "系统按样本类型和字段展示准确率统计，便于定位失败样本和易错字段。", "caption": "图2-2-1 意图解析参数提取准确率评测看板", "image": "intent-evaluation-dashboard.png"},
                {"text": "打开评测样本列表，查看样本输入、解析结果、期望结果、字段判定和失败字段。", "caption": "图2-2-2 意图解析评测样本明细", "image": "intent-evaluation-samples.png"},
            ],
            "expected": "参数提取准确率不小于 90%。",
            "result": "已在管理节点使用在线大模型/智能体对 360 条固定数据集完成逐条评测，正确 349 条，准确率 96.9%，达到 ≥90% 验收目标；评测编号为 online-eval-20260621-003559-2794dec2。",
        },
    ]

    for case in cases:
        add_case_table(doc, case)

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(OUT))
    print(OUT)


if __name__ == "__main__":
    build()
